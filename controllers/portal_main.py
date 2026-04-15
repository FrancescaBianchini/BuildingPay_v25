# -*- coding: utf-8 -*-
import base64
import logging
from datetime import date
from io import BytesIO
from odoo import http, _
from odoo.http import request
from odoo.addons.portal.controllers.portal import CustomerPortal, pager as portal_pager
from odoo.exceptions import AccessError, UserError

_logger = logging.getLogger(__name__)


class BuildingPayPortal(CustomerPortal):
    """
    Controller portale BuildingPay.
    Aggiunge sezioni:
    - Contratti: due documenti distinti per l'amministratore
        1. Accordo Retrocessioni Amministratore ED (download pre-compilato + upload firmato)
        2. Accordo Condomini Aggregati ED (download con Allegato A + upload firmato)
    - Condomini: CRUD indirizzi di tipo 'condominio'
    """

    # -------------------------------------------------------
    # Home portale: aggiunge le sezioni BuildingPay
    # -------------------------------------------------------
    def _prepare_home_portal_values(self, counters):
        values = super()._prepare_home_portal_values(counters)

        # ---------------------------------------------------------------
        # is_amministratore viene aggiunto SOLO sul rendering iniziale
        # della pagina (counters == []), MAI sulle chiamate AJAX di
        # aggiornamento contatori (counters == ['invoice_count', ...]).
        #
        # Motivo: il JS di Odoo 17 EE, alla risposta di /my/home/counts,
        # itera su TUTTE le chiavi del dict restituito e chiama:
        #   document.getElementById(key).textContent = value
        # Se una chiave (es. 'is_amministratore') non ha un corrispondente
        # <span id="is_amministratore"> nel DOM, getElementById restituisce
        # null → TypeError "Cannot set properties of null (setting 'textContent')".
        #
        # Sul rendering iniziale (counters=[]) il valore serve al template
        # per il t-if="is_amministratore". Sulle chiamate AJAX il template
        # NON viene ri-renderizzato, quindi il valore non serve.
        # ---------------------------------------------------------------
        if not counters:
            partner = request.env.user.partner_id.sudo()
            values['is_amministratore'] = partner.is_amministratore

        return values

    # -------------------------------------------------------
    # SEZIONE: Contratti (pagina principale)
    # -------------------------------------------------------
    @http.route('/my/contratti', type='http', auth='user', website=True)
    def portal_contratti(self, **kw):
        """Pagina Contratti nel portale: mostra i due documenti disponibili."""
        partner = request.env.user.partner_id
        config = request.env['buildingpay_v25.config'].sudo().get_config_for_website()

        # Conta condomini attivi per mostrare/nascondere l'Accordo Condomini
        condomini_count = request.env['res.partner'].sudo().search_count([
            ('parent_id', '=', partner.id),
            ('type', '=', 'condominio'),
            ('active', '=', True),
        ])

        values = {
            'partner': partner,
            'config': config,
            'condomini_count': condomini_count,
            'page_name': 'contratti',
            'success': kw.get('success'),
            'error': kw.get('error'),
        }
        return request.render('BuildingPay_v25.portal_contratti', values)

    # -------------------------------------------------------
    # CONTRATTO 1: Accordo Retrocessioni Amministratore ED
    # Download con placeholder: NOME AMMINISTRATORE, CODICE FISCALE, IBAN, NOME BANCA, DATA
    # -------------------------------------------------------
    @http.route('/my/contratti/retrocessioni/download', type='http', auth='user', website=True)
    def portal_retrocessioni_download(self, **kw):
        """
        Download del template 'Accordo Retrocessioni Amministratore ED'
        con sostituzione dei placeholder:
        - [NOME AMMINISTRATORE] → partner.name
        - [CODICE FISCALE]      → partner.fiscalcode
        - [IBAN]                → primo res.partner.bank.acc_number
        - [NOME BANCA]          → bank_id.name su res.partner.bank
        - [DATA]                → data odierna DD/MM/YYYY
        """
        partner = request.env.user.partner_id
        config = request.env['buildingpay_v25.config'].sudo().get_config_for_website()

        if not config or not config.accordo_retrocessioni_template:
            return request.redirect('/my/contratti?error=no_template_retrocessioni')

        try:
            from docx import Document

            template_data = base64.b64decode(config.accordo_retrocessioni_template)
            doc = Document(BytesIO(template_data))

            nome_amministratore = partner.name or ''
            codice_fiscale = partner.fiscalcode or ''

            # Recupera IBAN dal primo conto bancario del partner
            bank = request.env['res.partner.bank'].sudo().search([
                ('partner_id', '=', partner.id),
            ], limit=1)
            iban = bank.acc_number if bank else ''
            # Nome banca da res.partner.bank.bank_id (campo nativo Odoo)
            banca_nome = (bank.bank_id.name if bank and bank.bank_id else '') or ''
            oggi = date.today().strftime('%d/%m/%Y')

            replacements = {
                '[NOME AMMINISTRATORE]': nome_amministratore,
                '[CODICE FISCALE]': codice_fiscale,
                '[IBAN]': iban,
                '[NOME BANCA]': banca_nome,
                '[DATA]': oggi,
            }

            # Sostituzione nei paragrafi
            for paragraph in doc.paragraphs:
                for placeholder, value in replacements.items():
                    self._replace_placeholder_in_paragraph(paragraph, placeholder, value)

            # Sostituzione nelle tabelle
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in replacements.items():
                                self._replace_placeholder_in_paragraph(
                                    paragraph, placeholder, value)

            output = BytesIO()
            doc.save(output)
            output.seek(0)
            file_data = output.read()

            filename = 'Accordo Retrocessioni Amministratore ED.docx'
            return request.make_response(
                file_data,
                headers=[
                    ('Content-Type',
                     'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
                    ('Content-Disposition',
                     'attachment; filename="%s"' % filename),
                    ('Content-Length', len(file_data)),
                ],
            )

        except Exception as e:
            _logger.error('BuildingPay: errore download accordo retrocessioni: %s', e)
            return request.redirect('/my/contratti?error=download_error')

    @http.route('/my/contratti/retrocessioni/upload', type='http', auth='user',
                website=True, methods=['POST'])
    def portal_retrocessioni_upload(self, **kw):
        """
        Upload del file 'Accordo Retrocessioni' firmato.
        Attiva il flag accordo_retrocessioni_ed sul partner.
        """
        partner = request.env.user.partner_id
        uploaded_file = kw.get('retrocessioni_file')

        if not uploaded_file:
            return request.redirect('/my/contratti?error=no_file')

        try:
            file_data = uploaded_file.read()
            filename = uploaded_file.filename
            file_b64 = base64.b64encode(file_data)
            partner.sudo().action_upload_retrocessioni(file_b64, filename)
            return request.redirect('/my/contratti?success=retrocessioni_uploaded')
        except Exception as e:
            _logger.error('BuildingPay: errore upload accordo retrocessioni: %s', e)
            return request.redirect('/my/contratti?error=upload_error')

    # -------------------------------------------------------
    # CONTRATTO 2: Accordo Condomini Aggregati ED
    # Download con placeholder: NOME AMMINISTRATORE, [________] (CF), [ALLEGATO_A]
    # [ALLEGATO_A] viene sostituito da una tabella con i condomini attivi
    # -------------------------------------------------------
    @http.route('/my/contratti/condomini-aggregati/download', type='http',
                auth='user', website=True)
    def portal_contratto_download(self, **kw):
        """
        Download del template 'Accordo Condomini Aggregati ED'
        con sostituzione dei placeholder:
        - [NOME AMMINISTRATORE] → partner.name
        - [________]            → partner.fiscalcode
        - [ALLEGATO_A]          → tabella condomini attivi (nome | indirizzo | IBAN)
        """
        partner = request.env.user.partner_id
        config = request.env['buildingpay_v25.config'].sudo().get_config_for_website()

        if not config or not config.contratto_template:
            return request.redirect('/my/contratti?error=no_template_condomini')

        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            template_data = base64.b64decode(config.contratto_template)
            doc = Document(BytesIO(template_data))

            nome_amministratore = partner.name or ''
            codice_fiscale = partner.fiscalcode or ''

            # Recupera condomini attivi dell'amministratore
            condomini = request.env['res.partner'].sudo().search([
                ('parent_id', '=', partner.id),
                ('type', '=', 'condominio'),
                ('active', '=', True),
            ])

            # Sostituzione nei paragrafi (escluso [ALLEGATO_A] che richiede inserimento tabella)
            allegato_a_para = None
            for paragraph in doc.paragraphs:
                if '[ALLEGATO_A]' in paragraph.text:
                    allegato_a_para = paragraph
                    continue
                self._replace_placeholder_in_paragraph(
                    paragraph, '[NOME AMMINISTRATORE]', nome_amministratore)
                self._replace_placeholder_in_paragraph(
                    paragraph, '[________]', codice_fiscale)

            # Sostituzione nelle tabelle
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_placeholder_in_paragraph(
                                paragraph, '[NOME AMMINISTRATORE]', nome_amministratore)
                            self._replace_placeholder_in_paragraph(
                                paragraph, '[________]', codice_fiscale)

            # Inserimento tabella Allegato A
            if allegato_a_para is not None:
                self._insert_allegato_a_table(doc, allegato_a_para, condomini)

            output = BytesIO()
            doc.save(output)
            output.seek(0)
            file_data = output.read()

            filename = 'Accordo Condomini Aggregati ED.docx'
            return request.make_response(
                file_data,
                headers=[
                    ('Content-Type',
                     'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
                    ('Content-Disposition',
                     'attachment; filename="%s"' % filename),
                    ('Content-Length', len(file_data)),
                ],
            )

        except Exception as e:
            _logger.error('BuildingPay: errore download accordo condomini aggregati: %s', e)
            return request.redirect('/my/contratti?error=download_error')

    def _insert_allegato_a_table(self, doc, placeholder_paragraph, condomini):
        """
        Sostituisce il paragrafo [ALLEGATO_A] con una tabella Allegato A
        contenente: Denominazione | Indirizzo | IBAN
        """
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy

        # Prepara la tabella con intestazione
        headers = ['Denominazione Condominio', 'Indirizzo', 'IBAN']
        table = doc.add_table(rows=1 + len(condomini), cols=3)
        table.style = 'Table Grid'

        # Riga intestazione
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            run = hdr_cells[i].paragraphs[0].runs[0] if hdr_cells[i].paragraphs[0].runs else \
                hdr_cells[i].paragraphs[0].add_run(header)
            run.bold = True

        # Righe dati condomini
        for row_idx, condo in enumerate(condomini, 1):
            bank = request.env['res.partner.bank'].sudo().search([
                ('partner_id', '=', condo.id),
            ], limit=1)
            address_parts = filter(None, [
                condo.street, condo.zip, condo.city,
                condo.state_id.name if condo.state_id else '',
            ])
            address = ' '.join(address_parts)
            row_cells = table.rows[row_idx].cells
            row_cells[0].text = condo.name or ''
            row_cells[1].text = address
            row_cells[2].text = bank.acc_number if bank else ''

        # Sposta la tabella nella posizione del paragrafo placeholder
        # Usando l'API lxml: inserisce la tabella XML dopo il paragrafo placeholder
        para_element = placeholder_paragraph._p
        tbl_element = table._tbl
        para_element.addnext(tbl_element)

        # Svuota il paragrafo placeholder (lascia un paragrafo vuoto)
        for run in placeholder_paragraph.runs:
            run.text = ''
        if placeholder_paragraph.runs:
            pass
        else:
            placeholder_paragraph.clear()

    @http.route('/my/contratti/condomini-aggregati/upload', type='http', auth='user',
                website=True, methods=['POST'])
    def portal_contratto_upload(self, **kw):
        """
        Upload del file 'Accordo Condomini Aggregati ED' firmato.
        Attiva il flag accordo_condomini_aggregati_ed sul partner.
        """
        partner = request.env.user.partner_id
        uploaded_file = kw.get('contratto_file')

        if not uploaded_file:
            return request.redirect('/my/contratti?error=no_file')

        try:
            file_data = uploaded_file.read()
            filename = uploaded_file.filename
            file_b64 = base64.b64encode(file_data)
            partner.sudo().action_upload_accordo_condomini(file_b64, filename)
            return request.redirect('/my/contratti?success=condomini_uploaded')
        except Exception as e:
            _logger.error('BuildingPay: errore upload accordo condomini aggregati: %s', e)
            return request.redirect('/my/contratti?error=upload_error')

    # -------------------------------------------------------
    # SEZIONE: Condomini
    # -------------------------------------------------------
    @http.route('/my/condomini', type='http', auth='user', website=True)
    def portal_condomini_list(self, page=1, **kw):
        """Lista dei condomini dell'amministratore."""
        partner = request.env.user.partner_id

        if not partner.is_amministratore:
            return request.redirect('/my')

        domain = [
            ('parent_id', '=', partner.id),
            ('type', '=', 'condominio'),
            ('active', '=', True),
        ]

        condomini = request.env['res.partner'].sudo().search(domain)

        values = {
            'partner': partner,
            'condomini': condomini,
            'page_name': 'condomini',
        }
        return request.render('BuildingPay_v25.portal_condomini', values)

    @http.route('/my/condomini/new', type='http', auth='user', website=True)
    def portal_condominio_new(self, **kw):
        """Form per aggiungere un nuovo condominio."""
        partner = request.env.user.partner_id

        if not partner.is_amministratore:
            return request.redirect('/my')

        countries = request.env['res.country'].sudo().search([])
        values = {
            'partner': partner,
            'condominio': None,
            'countries': countries,
            'page_name': 'condomini_new',
            'mode': 'create',
        }
        return request.render('BuildingPay_v25.portal_condominio_form', values)

    @http.route('/my/condomini/new/save', type='http', auth='user', website=True,
                methods=['POST'])
    def portal_condominio_create(self, **kw):
        """Salva un nuovo indirizzo condominio."""
        partner = request.env.user.partner_id

        if not partner.is_amministratore:
            return request.redirect('/my')

        params = request.params
        errors = self._validate_condominio_form(params)

        if errors:
            countries = request.env['res.country'].sudo().search([])
            return request.render('BuildingPay_v25.portal_condominio_form', {
                'partner': partner,
                'condominio': None,
                'countries': countries,
                'errors': errors,
                'form_data': params,
                'mode': 'create',
                'page_name': 'condomini_new',
            })

        try:
            condominio_vals = self._prepare_condominio_vals(params, partner)
            condominio = request.env['res.partner'].sudo().create(condominio_vals)

            # Salva IBAN nei conti bancari
            iban = params.get('iban', '').strip()
            if iban:
                request.env['res.partner.bank'].sudo().create({
                    'partner_id': condominio.id,
                    'acc_number': iban,
                })

            # Attiva flag electronic invoice se codice destinatario presente
            if params.get('codice_destinatario'):
                condominio.sudo().write({
                    'electronic_invoice_subjected': True,
                    'electronic_invoice_obliged_subject': True,
                })

            return request.redirect('/my/condomini?success_add=1')
        except Exception as e:
            _logger.error('BuildingPay: errore creazione condominio: %s', e)
            return request.redirect('/my/condomini?error=create_error')

    @http.route('/my/condomini/<int:condominio_id>', type='http', auth='user', website=True)
    def portal_condominio_detail(self, condominio_id, **kw):
        """Dettaglio/modifica di un condominio esistente."""
        partner = request.env.user.partner_id
        condominio = self._get_condominio_or_redirect(condominio_id, partner)
        if isinstance(condominio, type(request.redirect('/'))):
            return condominio

        # Recupera IBAN dal conto bancario
        bank = request.env['res.partner.bank'].sudo().search([
            ('partner_id', '=', condominio.id),
        ], limit=1)

        countries = request.env['res.country'].sudo().search([])
        values = {
            'partner': partner,
            'condominio': condominio,
            'bank': bank,
            'countries': countries,
            'page_name': 'condomini_edit',
            'mode': 'edit',
        }
        return request.render('BuildingPay_v25.portal_condominio_form', values)

    @http.route('/my/condomini/<int:condominio_id>/save', type='http', auth='user',
                website=True, methods=['POST'])
    def portal_condominio_update(self, condominio_id, **kw):
        """Aggiorna un condominio esistente."""
        partner = request.env.user.partner_id
        condominio = self._get_condominio_or_redirect(condominio_id, partner)
        if isinstance(condominio, type(request.redirect('/'))):
            return condominio

        params = request.params
        errors = self._validate_condominio_form(params)

        if errors:
            countries = request.env['res.country'].sudo().search([])
            bank = request.env['res.partner.bank'].sudo().search([
                ('partner_id', '=', condominio.id),
            ], limit=1)
            return request.render('BuildingPay_v25.portal_condominio_form', {
                'partner': partner,
                'condominio': condominio,
                'bank': bank,
                'countries': countries,
                'errors': errors,
                'form_data': params,
                'mode': 'edit',
                'page_name': 'condomini_edit',
            })

        try:
            condominio_vals = self._prepare_condominio_vals(params, partner)
            # Non sovrascriviamo parent_id e type in aggiornamento
            condominio_vals.pop('parent_id', None)
            condominio_vals.pop('type', None)
            condominio.sudo().write(condominio_vals)

            # Aggiorna IBAN
            iban = params.get('iban', '').strip()
            existing_bank = request.env['res.partner.bank'].sudo().search([
                ('partner_id', '=', condominio.id),
            ], limit=1)
            if iban:
                if existing_bank:
                    existing_bank.sudo().write({'acc_number': iban})
                else:
                    request.env['res.partner.bank'].sudo().create({
                        'partner_id': condominio.id,
                        'acc_number': iban,
                    })

            # Aggiorna flag electronic invoice
            if params.get('codice_destinatario'):
                condominio.sudo().write({
                    'electronic_invoice_subjected': True,
                    'electronic_invoice_obliged_subject': True,
                })

            return request.redirect('/my/condomini?success_edit=1')
        except Exception as e:
            _logger.error('BuildingPay: errore aggiornamento condominio: %s', e)
            return request.redirect('/my/condomini?error=update_error')

    @http.route('/my/condomini/<int:condominio_id>/archive', type='http', auth='user',
                website=True, methods=['POST'])
    def portal_condominio_archive(self, condominio_id, **kw):
        """Archivia un condominio (lo rende non attivo)."""
        partner = request.env.user.partner_id
        condominio = self._get_condominio_or_redirect(condominio_id, partner)
        if isinstance(condominio, type(request.redirect('/'))):
            return condominio

        try:
            condominio.sudo().action_archive_condominio()
            return request.redirect('/my/condomini?success_archive=1')
        except Exception as e:
            _logger.error('BuildingPay: errore archiviazione condominio: %s', e)
            return request.redirect('/my/condomini?error=archive_error')

    # -------------------------------------------------------
    # Metodi di utilità
    # -------------------------------------------------------
    def _get_condominio_or_redirect(self, condominio_id, partner):
        """
        Verifica che il condominio esista e appartenga all'utente corrente.
        Ritorna il record condominio oppure un redirect.
        """
        condominio = request.env['res.partner'].sudo().browse(condominio_id)
        if (not condominio.exists() or
                condominio.parent_id.id != partner.id or
                condominio.type != 'condominio'):
            return request.redirect('/my/condomini')
        return condominio

    def _validate_condominio_form(self, params):
        """Valida i dati del form condominio. Ritorna dict degli errori."""
        errors = {}
        if not params.get('name', '').strip():
            errors['name'] = _('Il nome è obbligatorio.')
        if not params.get('street', '').strip():
            errors['street'] = _("L'indirizzo è obbligatorio.")
        if not params.get('city', '').strip():
            errors['city'] = _('La città è obbligatoria.')
        if not params.get('zip', '').strip():
            errors['zip'] = _('Il CAP è obbligatorio.')
        if not params.get('fiscalcode', '').strip():
            errors['fiscalcode'] = _('Il codice fiscale è obbligatorio.')
        return errors

    def _prepare_condominio_vals(self, params, parent_partner):
        """Prepara il dict dei valori per creare/aggiornare un condominio."""
        vals = {
            'name': params.get('name', '').strip(),
            'type': 'condominio',
            'parent_id': parent_partner.id,
            'street': params.get('street', '').strip(),
            'street2': params.get('street2', '').strip(),
            'city': params.get('city', '').strip(),
            'zip': params.get('zip', '').strip(),
            'fiscalcode': params.get('fiscalcode', '').strip(),
            'pec_mail': params.get('pec_mail', '').strip(),
            'codice_destinatario': params.get('codice_destinatario', '').strip(),
        }
        country_id = params.get('country_id')
        if country_id:
            vals['country_id'] = int(country_id)
        state_id = params.get('state_id')
        if state_id:
            vals['state_id'] = int(state_id)
        return vals

    def _replace_placeholder_in_paragraph(self, paragraph, placeholder, replacement):
        """
        Sostituisce un placeholder nel testo di un paragrafo docx,
        preservando la formattazione dei run.
        Il placeholder potrebbe essere distribuito su più run.
        """
        if placeholder not in paragraph.text:
            return

        full_text = ''.join(run.text for run in paragraph.runs)
        if placeholder not in full_text:
            return

        new_text = full_text.replace(placeholder, replacement)

        # Riscrivi: metti tutto il testo nel primo run e svuota gli altri
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ''
